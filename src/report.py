from __future__ import annotations

from pathlib import Path

from src.schema import NodeMeta


def render_meta_section(meta: NodeMeta) -> str:
    lines = ["", "### 附：来源与置信度", "", f"- 置信度：{meta.confidence}"]
    if meta.sources:
        lines.append("- 引用来源：")
        for source in meta.sources:
            tail = f"（{source.provider}）" if source.provider else ""
            url = f" {source.url}" if source.url else ""
            lines.append(f"  - {source.title}{tail}{url}")
    if meta.assumptions:
        lines.append("- 关键假设：")
        lines.extend(f"  - {item}" for item in meta.assumptions)
    if meta.missing_info:
        lines.append("- 信息缺口：")
        lines.extend(f"  - {item}" for item in meta.missing_info)
    if meta.risk_flags:
        lines.append("- 风险标记：")
        lines.extend(f"  - {item}" for item in meta.risk_flags)
    return "\n".join(lines) + "\n"


def write_markdown(markdown_text: str, out_dir: str | Path, name: str) -> Path:
    out = Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)
    path = out / f"{name}.md"
    path.write_text(markdown_text, encoding="utf-8")
    return path


def write_docx(markdown_text: str, out_dir: str | Path, name: str) -> Path:
    from docx import Document

    out = Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)
    path = out / f"{name}.docx"
    document = Document()
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            document.add_paragraph("")
            continue
        stripped = line.strip()
        if stripped.startswith("### "):
            document.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            document.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            indent = (len(line) - len(line.lstrip())) // 2
            document.add_paragraph(stripped[2:], style="List Bullet" if indent == 0 else "List Bullet 2")
        else:
            document.add_paragraph(stripped)
    document.save(str(path))
    return path


def write_node_report(markdown_text: str, out_dir: str | Path, name: str) -> dict[str, str]:
    md_path = write_markdown(markdown_text, out_dir, name)
    try:
        docx_path = write_docx(markdown_text, out_dir, name)
    except Exception as exc:  # pragma: no cover - optional dependency missing
        return {"markdown": str(md_path), "docx_error": f"{type(exc).__name__}: {exc}"}
    return {"markdown": str(md_path), "docx": str(docx_path)}
