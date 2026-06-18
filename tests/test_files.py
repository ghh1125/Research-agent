from __future__ import annotations

from pathlib import Path

from src.files import parse_file


def test_parse_docx(tmp_path: Path) -> None:
    from docx import Document

    document = Document()
    document.add_paragraph("公司简介：示例科技有限公司")
    document.add_paragraph("成立时间：2020 年")
    path = tmp_path / "bp.docx"
    document.save(str(path))

    result = parse_file(path)
    assert result.kind == "docx"
    assert "示例科技有限公司" in result.text
    assert result.error is None


def test_parse_xlsx_sheets(tmp_path: Path) -> None:
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "利润表"
    sheet.append(["项目", "2024", "2025"])
    sheet.append(["营业收入", 1000, 1500])
    sheet.append(["营业成本", 600, 800])
    path = tmp_path / "financials.xlsx"
    workbook.save(str(path))

    result = parse_file(path)
    assert result.kind == "xlsx"
    assert "利润表" in result.sheets
    assert result.sheets["利润表"][1][0] == "营业收入"


def test_parse_missing_file(tmp_path: Path) -> None:
    result = parse_file(tmp_path / "does_not_exist.pdf")
    assert result.kind == "missing"
    assert result.error is not None


def test_parse_unsupported_suffix(tmp_path: Path) -> None:
    path = tmp_path / "notes.csv"
    path.write_text("a,b,c", encoding="utf-8")
    result = parse_file(path)
    assert result.kind == "unsupported"
    assert result.error is not None
